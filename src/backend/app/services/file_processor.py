import pandas as pd
import openpyxl
from docx import Document
from typing import Dict, Any
import io
import base64


class FileProcessor:
    """Service for processing various file types"""

    async def process_document(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process a document file (DOCX, PDF)"""
        try:
            if filename.endswith('.docx'):
                return await self._process_docx(content)
            elif filename.endswith('.pdf'):
                return await self._process_pdf(content)
            else:
                raise ValueError(f"Unsupported document format: {filename}")
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")

    async def _process_docx(self, content: bytes) -> Dict[str, Any]:
        """Process DOCX file"""
        doc = Document(io.BytesIO(content))

        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        html = self._docx_to_html(doc)

        return {
            "content": "\n".join(full_text),
            "html": html,
            "paragraphs": len(doc.paragraphs),
            "tables": len(tables),
            "table_data": tables
        }

    def _docx_to_html(self, doc: Document) -> str:
        """Convert DOCX to HTML"""
        html_parts = ['<div class="document-content">']

        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                    html_parts.append(f'<h{level}>{para.text}</h{level}>')
                else:
                    runs_html = ""
                    for run in para.runs:
                        text = run.text
                        if run.bold:
                            text = f"<strong>{text}</strong>"
                        if run.italic:
                            text = f"<em>{text}</em>"
                        if run.underline:
                            text = f"<u>{text}</u>"
                        runs_html += text
                    html_parts.append(f'<p>{runs_html}</p>')

        for table in doc.tables:
            html_parts.append(
                '<table border="1" style="border-collapse: collapse; margin: 10px 0;">'
            )
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    html_parts.append(
                        f'<td style="padding: 5px; border: 1px solid #ccc;">'
                        f'{cell.text}</td>'
                    )
                html_parts.append('</tr>')
            html_parts.append('</table>')

        html_parts.append('</div>')
        return "".join(html_parts)

    async def _process_pdf(self, content: bytes) -> Dict[str, Any]:
        """Process Pdf file"""
        return {
            "content": "PDF content extraction not implemented yet",
            "html": "<p>PDF preview not available</p>",
            "paragraphs": 0,
            "tables": 0
        }

    async def process_spreadsheet(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process a spreadsheet file (XLSX, CSV)"""
        try:
            if filename.endswith('.csv'):
                return await self._process_csv(content)
            elif filename.endswith(('.xlsx', '.xls')):
                return await self._process_excel(content)
            else:
                raise ValueError(f"Unsupported spreadsheet format: {filename}")
        except Exception as e:
            raise Exception(f"Error processing spreadsheet: {str(e)}")

    async def _process_excel(self, content: bytes) -> Dict[str, Any]:
        """Process Excel file"""
        excel_file = io.BytesIO(content)
        workbook = openpyxl.load_workbook(excel_file)

        sheets = []
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]

            data = []
            headers = []

            for row_idx, row in enumerate(worksheet.iter_rows()):
                row_data = [cell.value for cell in row]
                if row_idx == 0:
                    headers = row_data
                else:
                    data.append(row_data)

            sheets.append({
                "name": sheet_name,
                "headers": headers,
                "rows": data,
                "row_count": len(data),
                "column_count": len(headers)
            })

        return {
            "sheets": sheets,
            "sheet_count": len(sheets)
        }

    async def _process_csv(self, content: bytes) -> Dict[str, Any]:
        """Process CSV file"""
        df = pd.read_csv(io.BytesIO(content))

        return {
            "sheets": [{
                "name": "Sheet1",
                "headers": df.columns.tolist(),
                "rows": df.values.tolist(),
                "row_count": len(df),
                "column_count": len(df.columns)
            }],
            "sheet_count": 1
        }

    async def export_to_excel(self, content: Dict[str, Any], filename: str) -> Dict[str, Any]:
        """Export data to Excel"""
        try:
            output = io.BytesIO()

            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                for sheet in content.get("sheets", []):
                    df = pd.DataFrame(
                        sheet.get("rows", []),
                        columns=sheet.get("headers", [])
                    )
                    df.to_excel(writer, sheet_name=sheet.get("name", "Sheet1"), index=False)

            output.seek(0)

            return {
                "success": True,
                "filename": filename,
                "data": base64.b64encode(output.getvalue()).decode()
            }
        except Exception as e:
            raise Exception(f"Error exporting to Excel: {str(e)}")

    async def export_to_csv(self, content: Dict[str, Any], filename: str) -> Dict[str, Any]:
        """Export data to CSV"""
        try:
            output = io.StringIO()

            sheet = content.get("sheets", [{}])[0]
            df = pd.DataFrame(
                sheet.get("rows", []),
                columns=sheet.get("headers", [])
            )
            df.to_csv(output, index=False)

            return {
                "success": True,
                "filename": filename,
                "data": base64.b64encode(output.getvalue().encode()).decode()
            }
        except Exception as e:
            raise Exception(f"Error exporting to CSV: {str(e)}")

    async def export_to_docx(self, content: Dict[str, Any], filename: str) -> Dict[str, Any]:
        """Export content to DOCX"""
        try:
            doc = Document()

            text_content = content.get("content", "")
            for paragraph in text_content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)

            for table_data in content.get("tables", []):
                rows = len(table_data)
                cols = len(table_data[0]) if table_data else 0
                table = doc.add_table(rows=rows, cols=cols)
                for i, row_data in enumerate(table_data):
                    for j, cell_data in enumerate(row_data):
                        table.cell(i, j).text = str(cell_data)

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            return {
                "success": True,
                "filename": filename,
                "data": base64.b64encode(output.getvalue()).decode()
            }
        except Exception as e:
            raise Exception(f"Error exporting to DOCX: {str(e)}")
